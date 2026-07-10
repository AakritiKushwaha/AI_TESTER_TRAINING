import sys, json, base64, os
from io import BytesIO
from datetime import datetime

try:
    from docx import Document
except ImportError:
    # If python-docx is not installed, raise a clear error.
    raise RuntimeError('python-docx library is required. Install with `pip install python-docx`.')

def generate_docx(sections, title=None):
    """Create a DOCX document from a list of sections.
    Each section is a dict with 'title' and 'content' (string, may contain markdown)."""
    doc = Document()
    # Optional title page
    if title:
        doc.add_heading(title, level=0)
        doc.add_paragraph(f"Generated on {datetime.utcnow().isoformat()} UTC")
        doc.add_page_break()
    for sec in sections:
        sec_title = sec.get('title', 'Section')
        content = sec.get('content', '')
        doc.add_heading(sec_title, level=1)
        # Simple handling: split by newlines for paragraphs
        for para in content.split('\n'):
            doc.add_paragraph(para)
    # Save to in‑memory buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()

def main(input_json: str):
    payload = json.loads(input_json)
    sections = payload.get('sections') or payload.get('result', {}).get('sections')
    if not sections:
        raise RuntimeError('No sections provided in payload')
    file_name = payload.get('docxFileName') or f"TestStrategy_{payload.get('jiraId', 'unknown')}.docx"
    docx_bytes = generate_docx(sections, title=file_name)
    b64 = base64.b64encode(docx_bytes).decode('utf-8')
    result = {
        'docxFileName': file_name,
        'docxBase64': b64,
    }
    print(json.dumps(result))

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print(json.dumps({"error": "Input JSON required"}))
        sys.exit(1)
    main(sys.argv[1])
